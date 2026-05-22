#!/usr/bin/env python3
"""
Generate PiggyWise Setup Guide (.docx) for each product in TH and EN.
Output: docs/instructions/th/  and  docs/instructions/en/
"""

import json
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

BASE_DIR   = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR    = os.path.join(BASE_DIR, "docs", "instructions")
LOGO_DIR   = os.path.join(BASE_DIR, "images", "logo")
COVER_DIR  = os.path.join(BASE_DIR, "images", "cover")
SLIDES_DIR = os.path.join(BASE_DIR, "images", "slides")
BANNER     = os.path.join(LOGO_DIR, "piggywise-banner.png")
ICON       = os.path.join(LOGO_DIR, "piggywise-icon.png")

COLOR_GOLD     = RGBColor(0xD4, 0xA0, 0x17)
COLOR_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_OFFWHITE = RGBColor(0xE8, 0xE0, 0xD0)
COLOR_DARK     = RGBColor(0x1A, 0x1A, 0x1A)
COLOR_GREY     = RGBColor(0x66, 0x66, 0x66)
COLOR_CODE_BG  = "F0F0F0"

with open(os.path.join(BASE_DIR, "data", "setup_guide.json"), encoding="utf-8") as _f:
    _data = json.load(_f)

PRODUCTS = _data["products"]


# ── Hyperlink helper ──────────────────────────────────────────────────────────

def add_hyperlink(para, text, url, color="D4A017", size=11, bold=False):
    r_id = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color)
    u = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz");   sz.set(qn("w:val"), str(int(size * 2)))
    rPr.append(c); rPr.append(u); rPr.append(sz)
    if bold:
        b = OxmlElement("w:b"); rPr.append(b)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


# ── XML helpers ────────────────────────────────────────────────────────────────

def shading(element, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    return shd


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(shading(tcPr, hex_color))


def set_cell_borders(cell, color="D4A017", size=18):
    tcPr = cell._tc.get_or_add_tcPr()
    bdr = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        bdr.append(b)
    tcPr.append(bdr)


def set_para_shading(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    pPr.append(shading(pPr, hex_color))


def insert_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    p.add_run()._r.append(br)


def add_bottom_border(para, color="D4A017"):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


# ── Footer ─────────────────────────────────────────────────────────────────────

def add_footer(doc, short_name):
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    footer = section.footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    tbl.style = "Table Grid"

    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    tbl.columns[0].width = Inches(2.8)
    tbl.columns[1].width = Inches(0.9)
    tbl.columns[2].width = Inches(2.8)

    left = tbl.cell(0, 0)
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(ICON):
        lp.add_run().add_picture(ICON, height=Pt(12))
    r = lp.add_run(f"  PiggyWise  ·  {short_name}")
    r.font.size = Pt(8)
    r.font.color.rgb = COLOR_GREY

    tbl.cell(0, 1).paragraphs[0].text = ""

    right = tbl.cell(0, 2)
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r_pre = rp.add_run("p. ")
    r_pre.font.size = Pt(8)
    r_pre.font.color.rgb = COLOR_GREY

    for field, label in [("PAGE", None), (None, " / "), ("NUMPAGES", None)]:
        if field:
            fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin")
            ins = OxmlElement("w:instrText"); ins.text = field
            fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate")
            fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end")
            run = rp.add_run()
            run._r.extend([fc1, ins, fc2, fc3])
            run.font.size = Pt(8)
            run.font.color.rgb = COLOR_GREY
        else:
            r = rp.add_run(label)
            r.font.size = Pt(8)
            r.font.color.rgb = COLOR_GREY


# ── Cover page ─────────────────────────────────────────────────────────────────

def _add_banner_inline(doc):
    if os.path.exists(BANNER):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(BANNER, width=Inches(6.3))


def add_cover_page(doc, product_key, lang, lang_data):
    cover_img  = os.path.join(COVER_DIR, f"cover_{product_key}_{lang}.png")
    slide1_img = os.path.join(SLIDES_DIR, product_key, f"slide1_features_{lang}.png")

    # Page 1: banner + cover image
    _add_banner_inline(doc)
    if os.path.exists(cover_img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(cover_img, width=Inches(6.3))

    insert_page_break(doc)

    # Page 2: banner + thank-you + slide 1
    _add_banner_inline(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(lang_data["thank_you"])
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = COLOR_GOLD

    if os.path.exists(slide1_img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(slide1_img, width=Inches(6.3))

    insert_page_break(doc)


# ── Content helpers ────────────────────────────────────────────────────────────

def add_banner(doc):
    if os.path.exists(BANNER):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        p.add_run().add_picture(BANNER, width=Inches(6.3))


def add_title(doc, product_name, guide_title):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(guide_title)
    r1.font.size = Pt(26)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_DARK

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(product_name)
    r2.font.size = Pt(13)
    r2.font.bold = True
    r2.font.color.rgb = COLOR_GOLD
    add_bottom_border(p2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = COLOR_DARK
    add_bottom_border(p)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        r.font.size = Pt(11)
        r.font.color.rgb = COLOR_DARK
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_DARK
    return p


def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{num}. {text}")
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_DARK
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"ℹ️  {text}")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = COLOR_GREY
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    set_para_shading(p, COLOR_CODE_BG)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_DARK
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = COLOR_DARK
    return p


# ── Product card ───────────────────────────────────────────────────────────────

def add_product_card(doc, product_name, product_link, strings):
    insert_page_break(doc)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "0D0D0D")
    set_cell_borders(cell, "D4A017", 18)

    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    def cp(text, size, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
           space_before=0, space_after=8, italic=False, font=None, underline=False):
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.underline = underline
        r.font.color.rgb = color or COLOR_WHITE
        if font:
            r.font.name = font
        return p

    cp(strings["product_card_label"], 11, bold=True, color=COLOR_GOLD, space_before=36, space_after=4)
    cp(product_name, 15, bold=True, color=COLOR_WHITE, space_after=8)
    cp("─" * 42, 9, color=COLOR_GOLD, space_after=12)
    cp(strings["product_card_link_label"], 10, color=COLOR_OFFWHITE, space_after=6)
    cp(product_link, 12, bold=True, color=COLOR_GOLD, font="Courier New", underline=True, space_after=16)
    cp(strings["product_card_instruction"], 9, italic=True, color=COLOR_OFFWHITE, space_after=28)


# ── Tab overview table ─────────────────────────────────────────────────────────

def add_tab_table(doc, rows, headers):
    tbl = doc.add_table(rows=1 + len(rows), cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        set_cell_bg(c, "0D0D0D")
        p = c.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_GOLD

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.cell(ri + 1, ci)
            p = c.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK


# ── Ongoing use table ──────────────────────────────────────────────────────────

def add_ongoing_table(doc, rows, headers):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = tbl.cell(0, i)
        set_cell_bg(c, "0D0D0D")
        p = c.paragraphs[0]
        p.clear()
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_GOLD

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.cell(ri + 1, ci)
            p = c.paragraphs[0]
            p.clear()
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ── Document builder ───────────────────────────────────────────────────────────

def build_doc(product_name, short_name, product_key, product_link, lang, lang_data):
    s   = lang_data["strings"]
    act = lang_data["activation_cmd"]

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Tahoma"
    style.font.size = Pt(11)

    add_footer(doc, short_name)
    add_cover_page(doc, product_key, lang, lang_data)
    add_banner(doc)
    add_title(doc, product_name, s["guide_title"])

    # ── Overview ───────────────────────────────────────────────────────────────
    add_h1(doc, s["section_overview"])
    add_body(doc, s["overview_body"].format(product_name=product_name))
    doc.add_paragraph()
    add_body(doc, s["tab_color_system"])
    for line in s["tab_color_bullets"]:
        add_bullet(doc, line)
    doc.add_paragraph()
    add_body(doc, s["tab_all_tabs"])
    add_tab_table(doc, [tuple(r) for r in lang_data["tabs"]], s["tab_headers"])

    # ── Product card ───────────────────────────────────────────────────────────
    add_product_card(doc, product_name, product_link, s)

    # ── Step 1 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step1_title"])
    add_body(doc, s["step1_body"])
    add_note(doc, s["step1_note"])
    doc.add_paragraph()

    # ── Step 2 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step2_title"])
    add_body(doc, s["step2_body"])
    for item in s["step2_bullets"]:
        add_bullet(doc, item)
    doc.add_paragraph()

    # ── Step 3 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step3_title"])
    add_body(doc, s["step3_body"])
    for item in s["step3_bullets"]:
        add_bullet(doc, item)
    add_note(doc, s["step3_note"])
    doc.add_paragraph()

    # ── Step 4 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step4_title"])
    add_body(doc, s["step4_body"])
    for item in s["step4_bullets"]:
        add_bullet(doc, item)
    add_note(doc, s["step4_note"])
    doc.add_paragraph()

    # ── Step 5 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step5_title"])
    for i, step in enumerate(s["step5_numbered"], 1):
        add_numbered(doc, step, i)
    add_code(doc, act)
    add_body(doc, s["step5_body"])
    doc.add_paragraph()

    # ── Step 6 ─────────────────────────────────────────────────────────────────
    add_h1(doc, s["step6_title"])
    add_body(doc, s["step6_body"])
    doc.add_paragraph()

    add_subheading(doc, s["step6_method1_heading"])
    add_code(doc, s["step6_method1_cmd"])
    add_body(doc, s["step6_example_label"])
    add_code(doc, s["step6_method1_example"].format(product_name=product_name))

    add_subheading(doc, s["step6_method2_heading"])
    add_body(doc, s["step6_method2_body"])
    add_code(doc, s["step6_method2_cmd"])
    doc.add_paragraph()

    # ── Ongoing use ────────────────────────────────────────────────────────────
    add_h1(doc, s["ongoing_title"])
    add_ongoing_table(doc, [tuple(r) for r in lang_data["ongoing"]], s["ongoing_headers"])

    # ── Troubleshooting ────────────────────────────────────────────────────────
    add_h1(doc, s["troubles_title"])
    for i, t in enumerate(lang_data["troubles"], 1):
        cmd = t["cmd"].format(activation_cmd=act) if t["cmd"] else None
        add_subheading(doc, f"{i}. {t['title']}")
        add_body(doc, t["desc"])
        for step in t["steps"]:
            add_bullet(doc, step)
        if cmd:
            add_code(doc, cmd)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Test commands ──────────────────────────────────────────────────────────
    add_h1(doc, s["tests_title"])
    add_body(doc, s["tests_body"])
    doc.add_paragraph()
    for i, t in enumerate(lang_data["tests"], 1):
        cmd = t["cmd"].format(activation_cmd=act, product_name=product_name)
        add_subheading(doc, f"{i}. {t['title']}")
        add_body(doc, t["desc"])
        add_code(doc, cmd)
        add_note(doc, t["check"])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Wealth advisory ────────────────────────────────────────────────────────
    add_h1(doc, s["advisories_title"])
    add_body(doc, s["advisories_body"])
    doc.add_paragraph()
    for i, a in enumerate(lang_data["advisories"], 1):
        cmd = a["cmd"].format(product_name=product_name)
        add_subheading(doc, f"{i}. {a['title']}")
        add_body(doc, a["desc"])
        add_code(doc, cmd)
        add_note(doc, a["note"])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Beyond portfolio ───────────────────────────────────────────────────────
    add_h1(doc, s["beyond_title"])
    doc.add_paragraph()
    for i, b in enumerate(lang_data["beyond"], 1):
        add_subheading(doc, f"{i}. {b['title']}")
        add_code(doc, b["cmd"])
        add_note(doc, b["note"])
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Crisis & Risk ──────────────────────────────────────────────────────────
    add_h1(doc, s["crisis_title"])
    add_body(doc, s["crisis_body1"])
    add_body(doc, s["crisis_body2"])
    add_code(doc, s["crisis_cmd"])
    doc.add_paragraph()

    # ── Lifestyle ──────────────────────────────────────────────────────────────
    add_h1(doc, s["lifestyle_title"])
    add_body(doc, s["lifestyle_body1"])
    add_body(doc, s["lifestyle_body2"])
    add_code(doc, s["lifestyle_cmd"])
    doc.add_paragraph()

    # ── Executive summaries ────────────────────────────────────────────────────
    add_h1(doc, s["exec_title"])
    add_body(doc, s["exec_body1"])
    add_body(doc, s["exec_body2"])
    add_code(doc, s["exec_cmd"])
    doc.add_paragraph()

    # ── Closing ────────────────────────────────────────────────────────────────
    insert_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run(lang_data["closing_text"])
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = COLOR_GOLD

    # ── Resources ──────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_h1(doc, s["resources_title"])

    def add_resource_row(label, url_th, url_en):
        url = url_th if lang == "th" else url_en
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        rl = p.add_run(f"{label}   ")
        rl.font.size = Pt(11)
        rl.font.bold = True
        rl.font.color.rgb = COLOR_DARK
        add_hyperlink(p, "→ อ่านที่นี่" if lang == "th" else "→ Read here", url)

    add_resource_row(s["faq_label"],        s["faq_url_th"],        s["faq_url_en"])
    add_resource_row(s["disclaimer_label"], s["disclaimer_url_th"], s["disclaimer_url_en"])
    add_resource_row(s["policies_label"],   s["policies_url_th"],   s["policies_url_en"])

    # ── Contact ────────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_h1(doc, s["contact_title"])
    add_body(doc, s["contact_body"])

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    smoke = "--smoke" in sys.argv
    lang_arg = next((a.split("=")[1] for a in sys.argv if a.startswith("--lang=")), None)
    langs = [lang_arg] if lang_arg else ["th", "en"]

    for lang in langs:
        lang_data = _data[lang]
        out_dir = os.path.join(OUT_DIR, lang)
        os.makedirs(out_dir, exist_ok=True)
        products = PRODUCTS[:1] if smoke else PRODUCTS
        label = f"smoke ({lang})" if smoke else lang
        print(f"Generating PiggyWise Setup Guides — {label}...")
        for p in products:
            name  = p["name"][lang]
            link  = p["link"][lang]
            doc   = build_doc(name, p["short_name"], p["key"], link, lang, lang_data)
            safe  = p["short_name"].replace(" ", "_")
            out   = os.path.join(out_dir, f"PiggyWise_Setup_Guide_{safe}.docx")
            doc.save(out)
            print(f"  ✓ {out}")

    print(f"\nDone.")
